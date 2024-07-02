import os
import logging
import pymupdf
import docx
import hashlib
import pandas as pd
import uuid
import asyncio
from typing import List, Iterator, Set, Optional
import sys
from dotenv import load_dotenv

# Ensure the base directory is in the Python path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')))

from vectordbs.factory import get_datastore
from vectordbs.vector_store import VectorStore
from vectordbs.data_types import Document, DocumentChunk
from vectordbs.utils.watsonx import get_embeddings

# Import chunking methods
from chunking import simple_chunking, semantic_chunking

load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Define constants
DATA_DIR: str = os.environ.get("DATA_DIR", '/Users/mg/mg-work/manav/work/ai-experiments/rag_modulo/data')
VECTOR_DB: str = os.environ.get("VECTOR_DB", 'milvus')
MAX_WORKERS: int = 4  # Number of threads for concurrent processing
COLLECTION_NAME: str = os.environ.get("COLLECTION_NAME", 'default_collection')

# Chunking strategy and parameters from .env
CHUNKING_STRATEGY: str = os.getenv("CHUNKING_STRATEGY", "fixed")
CHUNK_SIZE: int = int(os.getenv("CHUNK_SIZE", "1000"))
CHUNK_OVERLAP: int = int(os.getenv("CHUNK_OVERLAP", "100"))
SEMANTIC_THRESHOLD: float = float(os.getenv("SEMANTIC_THRESHOLD", "0.8"))

def get_chunking_method():
    """
    Get the appropriate chunking method based on the CHUNKING_STRATEGY environment variable.
    """
    if CHUNKING_STRATEGY.lower() == "semantic":
        return lambda text: semantic_chunking(text, CHUNK_SIZE, SEMANTIC_THRESHOLD)
    else:
        return lambda text: simple_chunking(text, CHUNK_SIZE, CHUNK_OVERLAP)

def extract_text_from_page(page: pymupdf.Page) -> str:
    """
    Extract text content from a PDF page.

    Args:
        page (pymupdf.Page): The PDF page object.

    Returns:
        str: Extracted text from the page.
    """
    return page.get_text()

def extract_tables_from_page(page: pymupdf.Page) -> List[List[List[str]]]:
    """
    Extract tables from a PDF page.

    Args:
        page (pymupdf.Page): The PDF page object.

    Returns:
        List[List[List[str]]]: A list of extracted tables, where each table is a list of rows,
        and each row is a list of cell values.
    """
    tables = page.find_tables()
    extracted_tables: List[List[List[str]]] = []

    for table in tables:
        extracted_table = table.extract()
        cleaned_table = [[clean_text(cell) if cell is not None else '' for cell in row] for row in extracted_table]
        if any(cell for row in cleaned_table for cell in row):  # Check if table has any non-empty cells
            extracted_tables.append(cleaned_table)

    return extracted_tables

def clean_text(text: Optional[str]) -> str:
    """
    Clean and normalize text by removing special characters and extra whitespace.

    Args:
        text (Optional[str]): The input text to clean.

    Returns:
        str: The cleaned and normalized text.
    """
    if text is None:
        return ''
    # Remove special characters and extra whitespace
    cleaned_text = ''.join(char for char in text if char.isalnum() or char.isspace())
    cleaned_text = ' '.join(cleaned_text.split())
    return cleaned_text

def extract_images_from_page(page: pymupdf.Page, output_folder: str, saved_image_hashes: Set[str]) -> List[str]:
    """
    Extract images from a PDF page and save them to the output folder.

    Args:
        page (pymupdf.Page): The PDF page object.
        output_folder (str): The folder to save extracted images.
        saved_image_hashes (Set[str]): A set of hashes of already saved images to avoid duplicates.

    Returns:
        List[str]: A list of paths to the saved image files.
    """
    image_list = page.get_images(full=True)
    images: List[str] = []

    for img_index, img in enumerate(image_list, start=1):
        xref = img[0]
        try:
            base_image = page.parent.extract_image(xref)
            if base_image:
                image_bytes = base_image["image"]
                image_hash = hashlib.md5(image_bytes).hexdigest()
                logger.info(f"Processing image on page {page.number + 1}, index {img_index}, hash: {image_hash}")

                if image_hash not in saved_image_hashes:
                    image_extension = base_image["ext"]
                    image_filename = f"{output_folder}/image_{page.number + 1}_{img_index}.{image_extension}"
                    with open(image_filename, "wb") as img_file:
                        img_file.write(image_bytes)
                    images.append(image_filename)
                    saved_image_hashes.add(image_hash)
                    logger.info(f"Saved new image: {image_filename}")
                else:
                    logger.info(f"Skipped duplicate image on page {page.number + 1}, image index {img_index}")
            else:
                logger.warning(f"Failed to extract image {xref} from page {page.number + 1}")
        except Exception as e:
            logger.error(f"Error extracting image {xref} from page {page.number + 1}: {e}")

    return images

def read_txt(file_path: str) -> Iterator[Document]:
    """
    Read a text file and yield Document objects.

    Args:
        file_path (str): The path to the text file.

    Yields:
        Document: Document objects containing chunks of the text file.
    """
    chunking_method = get_chunking_method()

    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
            chunks = chunking_method(text)

            for chunk in chunks:
                document = Document(
                    name=os.path.basename(file_path),
                    document_id=str(uuid.uuid4()),
                    chunks=[DocumentChunk(chunk_id=str(uuid.uuid4()), text=chunk)]
                )
                yield document
    except Exception as e:
        logger.error(f"Error reading TXT file {file_path}: {e}")

def extract_metadata(file_path: str) -> Dict[str, Any]:
    """
    Extract metadata from the file.

    Args:
        file_path (str): The path to the file.

    Returns:
        Dict[str, Any]: A dictionary containing file metadata.
    """
    try:
        file_stat = os.stat(file_path)
        return {
            "file_name": os.path.basename(file_path),
            "file_path": file_path,
            "file_size": file_stat.st_size,
            "creation_time": file_stat.st_ctime,
            "modification_time": file_stat.st_mtime,
            "file_extension": os.path.splitext(file_path)[1].lower(),
        }
    except Exception as e:
        logger.error(f"Error extracting metadata for file {file_path}: {str(e)}")
        return {}

def read_pdf(file_path: str, output_folder: str, saved_image_hashes: Set[str]) -> Iterator[Document]:
    """
    Read a PDF file and yield Document objects for text, tables, and images.

    Args:
        file_path (str): The path to the PDF file.
        output_folder (str): The folder to save extracted images.
        saved_image_hashes (Set[str]): A set of hashes of already saved images to avoid duplicates.

    Yields:
        Document: Document objects containing chunks of the PDF content.
    """
    chunking_method = get_chunking_method()

    try:
        with pymupdf.open(file_path) as doc:
            for page in doc:
                text = extract_text_from_page(page)
                chunks = chunking_method(text)
                for chunk in chunks:
                    document = get_document(os.path.basename(file_path),
                                            str(uuid.uuid4()), chunk)
                    yield document

                tables = extract_tables_from_page(page)
                for table in tables:
                    table_text = "\n".join([" | ".join(row) for row in table])
                    if table_text:
                        table_chunks = chunking_method(table_text)
                        for table_text in table_chunks:
                            document = get_document(os.path.basename(file_path),
                                                    str(uuid.uuid4()), table_text)
                            yield document

                extract_images_from_page(page, output_folder, saved_image_hashes)
    except Exception as e:
        logger.error(f"Error reading PDF file {file_path}: {e}")
        raise

def get_document(name: str, document_id: str, text: str) -> Document:
    """
    Create a Document object with embedded vectors.

    Args:
        name (str): The name of the document.
        document_id (str): The unique identifier for the document.
        text (str): The text content of the document.

    Returns:
        Document: A Document object with embedded vectors.
    """
    return Document(
        name=name,
        document_id=document_id,
        chunks=[DocumentChunk(chunk_id=str(uuid.uuid4()), text=text,
                              vectors=get_embeddings(text),
                              document_id=document_id
                              )]
    )

def read_word(file_path: str) -> str:
    """
    Read a Word document and extract its text content, including headers, footers, and tables.

    Args:
        file_path (str): The path to the Word document.

    Returns:
        str: The extracted text content of the Word document.
    """
    try:
        doc = docx.Document(file_path)
        full_text = []

        # Extract headers
        for section in doc.sections:
            for header in section.header.paragraphs:
                full_text.append(header.text)

        # Extract main content
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))

        # Extract footers
        for section in doc.sections:
            for footer in section.footer.paragraphs:
                full_text.append(footer.text)

        return "\n".join(full_text)
    except docx.opc.exceptions.PackageNotFoundError:
        logger.error(f"File not found or not a valid Word document: {file_path}")
        return ""
    except Exception as e:
        logger.error(f"Error reading Word file {file_path}: {str(e)}")
        return ""


def read_excel(file_path: str) -> str:
    """
    Read an Excel file and convert its content to a string representation,
    handling multiple sheets and preserving sheet names.

    Args:
        file_path (str): The path to the Excel file.

    Returns:
        str: A string representation of the Excel file's content.
    """
    try:
        # Try reading with pandas first
        sheets_data = pd.read_excel(file_path, sheet_name=None)
        full_text = []

        for sheet_name, df in sheets_data.items():
            full_text.append(f"Sheet: {sheet_name}")
            full_text.append(df.to_string(index=False))
            full_text.append("\n")

        return "\n".join(full_text)
    except Exception as pandas_error:
        logger.warning(f"Pandas failed to read Excel file {file_path}: {str(pandas_error)}. Trying with openpyxl.")

        try:
            # Fallback to openpyxl for more complex Excel files
            wb = load_workbook(file_path, read_only=True, data_only=True)
            full_text = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                full_text.append(f"Sheet: {sheet_name}")

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                    full_text.append(row_text)

                full_text.append("\n")

            wb.close()
            return "\n".join(full_text)
        except InvalidFileException:
            logger.error(f"Invalid Excel file: {file_path}")
            return ""
        except Exception as e:
            logger.error(f"Error reading Excel file {file_path}: {str(e)}")
            return ""

async def process_and_store_document(document: Document, vector_store: VectorStore, collection_name: str) -> None:
    """
    Process and store a document in the vector store.

    Args:
        document (Document): The document to process and store.
        vector_store (VectorStore): The vector store to use for storage.
        collection_name (str): The name of the collection to store the document in.
    """
    try:
        await vector_store.add_documents_async(collection_name, [document])
        logger.info(f"Document {document.document_id} successfully stored in collection {collection_name}.")
    except Exception as e:
        logger.error(f"Error processing document {document.document_id}: {e}")
        raise

async def ingest_documents(data_dir: str, vector_store: VectorStore) -> None:
    """
    Ingest documents from the specified directory into the vector store.

    Args:
        data_dir (str): The directory containing the documents to ingest.
        vector_store (VectorStore): The vector store to use for storage.
    """
    output_folder = os.path.join(data_dir, 'extracted_images')
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    saved_image_hashes: Set[str] = set()

    tasks = []
    for root, _, files in os.walk(data_dir):
        for file in files:
            file_path = os.path.join(root, file)
            if file.endswith('.txt'):
                for document in read_txt(file_path):
                    tasks.append(process_and_store_document(document, vector_store, COLLECTION_NAME))
            elif file.endswith('.pdf'):
                for document in read_pdf(file_path, output_folder, saved_image_hashes):
                    tasks.append(process_and_store_document(document, vector_store, COLLECTION_NAME))
            elif file.endswith('.docx'):
                text = read_word(file_path)
                document = get_document(os.path.basename(file_path), str(uuid.uuid4()), text)
                tasks.append(process_and_store_document(document, vector_store, COLLECTION_NAME))
            elif file.endswith('.xlsx'):
                text = read_excel(file_path)
                document = get_document(os.path.basename(file_path), str(uuid.uuid4()), text)
                tasks.append(process_and_store_document(document, vector_store, COLLECTION_NAME))
            else:
                logger.warning(f"Unsupported file type: {file}")

    await asyncio.gather(*tasks)

async def main() -> None:
    """
    Main function to orchestrate the document ingestion process.
    """
    try:
        vector_store = get_datastore(VECTOR_DB)
        await vector_store.delete_collection_async(COLLECTION_NAME)
        await vector_store.create_collection_async(COLLECTION_NAME)
        await ingest_documents(DATA_DIR, vector_store)
        logger.info("Ingestion process completed successfully.")
    except Exception as e:
        logger.error(f"Error in main function: {e}")
        raise

if __name__ == "__main__":
    asyncio.run(main())
